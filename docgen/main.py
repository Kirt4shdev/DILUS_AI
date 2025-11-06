from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import List, Dict, Any
from datetime import datetime

app = FastAPI(
    title="DILUS_AI DocGen",
    description="Servicio de generación de documentos DOCX",
    version="2.0.0"
)

# ============================================
# MODELOS
# ============================================

class OfertaData(BaseModel):
    cliente: str
    proyecto: str
    propuesta_tecnica: str
    alcance: str
    plazos: str
    conceptos_precio: List[str]

class DocumentacionData(BaseModel):
    titulo: str
    tipo_documento: str
    contenido: str
    secciones: List[Dict[str, Any]]

# ============================================
# UTILIDADES
# ============================================

def add_styled_paragraph(doc, text, style='Normal', bold=False, size=None, align=None):
    """Añadir párrafo con estilo"""
    p = doc.add_paragraph(text, style=style)
    
    if bold:
        for run in p.runs:
            run.bold = True
    
    if size:
        for run in p.runs:
            run.font.size = Pt(size)
    
    if align:
        p.alignment = align
    
    return p

def add_header(doc, text, level=1):
    """Añadir encabezado con estilo"""
    heading = doc.add_heading(text, level=level)
    
    # Estilo personalizado
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(37, 99, 235)  # Azul corporativo
            run.font.size = Pt(18)
    
    return heading

def add_footer(doc, text):
    """Añadir pie de página"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

# ============================================
# ENDPOINTS
# ============================================

@app.get("/health")
def health():
    """Health check"""
    return {
        "status": "ok",
        "service": "DILUS_AI DocGen",
        "version": "2.0.0",
        "timestamp": datetime.now().isoformat()
    }

@app.post("/generate/oferta")
async def generate_oferta(data: OfertaData):
    """
    Generar documento DOCX de oferta comercial
    """
    try:
        # Crear documento
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título principal
        add_header(doc, 'PROPUESTA TÉCNICA Y COMERCIAL', level=1)
        
        # Fecha
        add_styled_paragraph(
            doc,
            f'Fecha: {datetime.now().strftime("%d de %B de %Y")}',
            align=WD_ALIGN_PARAGRAPH.RIGHT
        )
        doc.add_paragraph()
        
        # Cliente
        add_header(doc, 'Cliente', level=2)
        add_styled_paragraph(doc, data.cliente, bold=True, size=12)
        doc.add_paragraph()
        
        # Proyecto
        add_header(doc, 'Proyecto', level=2)
        add_styled_paragraph(doc, data.proyecto)
        doc.add_paragraph()
        
        # Propuesta Técnica
        add_header(doc, 'Propuesta Técnica', level=2)
        add_styled_paragraph(doc, data.propuesta_tecnica)
        doc.add_paragraph()
        
        # Alcance del Proyecto
        add_header(doc, 'Alcance del Proyecto', level=2)
        add_styled_paragraph(doc, data.alcance)
        doc.add_paragraph()
        
        # Plazos de Ejecución
        add_header(doc, 'Plazos de Ejecución', level=2)
        add_styled_paragraph(doc, data.plazos)
        doc.add_paragraph()
        
        # Estructura de Precios
        add_header(doc, 'Estructura de Precios', level=2)
        for concepto in data.conceptos_precio:
            doc.add_paragraph(concepto, style='List Bullet')
        
        doc.add_paragraph()
        
        # Nota final
        p = doc.add_paragraph()
        p.add_run('Nota: ').bold = True
        p.add_run('Esta propuesta tiene una validez de 30 días a partir de la fecha de emisión.')
        
        # Footer
        add_footer(doc, 'DILUS_AI - Soluciones de Ingeniería con IA | www.dilusai.com')
        
        # Guardar en memoria
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Retornar como stream
        filename = f"oferta_{data.cliente.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar oferta: {str(e)}")

@app.post("/generate/documentacion")
async def generate_documentacion(data: DocumentacionData):
    """
    Generar documentación técnica en formato DOCX
    """
    try:
        # Crear documento
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título principal
        add_header(doc, data.titulo, level=1)
        
        # Tipo de documento y fecha
        add_styled_paragraph(
            doc,
            f'{data.tipo_documento} | {datetime.now().strftime("%d de %B de %Y")}',
            align=WD_ALIGN_PARAGRAPH.CENTER
        )
        doc.add_paragraph()
        
        # Línea separadora
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Contenido principal
        if data.contenido:
            add_header(doc, 'Introducción', level=2)
            add_styled_paragraph(doc, data.contenido)
            doc.add_paragraph()
        
        # Secciones
        for i, seccion in enumerate(data.secciones, 1):
            titulo = seccion.get('titulo', f'Sección {i}')
            contenido = seccion.get('contenido', '')
            
            add_header(doc, f'{i}. {titulo}', level=2)
            add_styled_paragraph(doc, contenido)
            doc.add_paragraph()
        
        # Página de cierre
        doc.add_page_break()
        add_header(doc, 'Fin del Documento', level=3)
        add_styled_paragraph(
            doc,
            f'Generado automáticamente por DILUS_AI el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}',
            align=WD_ALIGN_PARAGRAPH.CENTER
        )
        
        # Footer
        add_footer(doc, 'DILUS_AI - Documentación Técnica Inteligente | www.dilusai.com')
        
        # Guardar en memoria
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Retornar como stream
        filename = f"{data.tipo_documento.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar documentación: {str(e)}")

# ============================================
# STARTUP MESSAGE
# ============================================

@app.on_event("startup")
async def startup_event():
    print("=" * 60)
    print("🚀 DILUS_AI DocGen Service Started")
    print("=" * 60)
    print("📄 Service: Document Generation (DOCX)")
    print("🔗 Endpoints:")
    print("   - POST /generate/oferta")
    print("   - POST /generate/documentacion")
    print("   - GET  /health")
    print("=" * 60)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8090)

